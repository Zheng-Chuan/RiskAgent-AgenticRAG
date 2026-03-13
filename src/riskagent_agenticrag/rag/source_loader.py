"""语料加载.

支持从 corpus 目录递归读取多种格式文件：
- Markdown (.md)
- PDF (.pdf)
- Word文档 (.docx)
- Excel表格 (.xlsx, .xls)
- HTML网页 (.html, .htm)

约定.
- 读取支持的文件格式
- 跳过 README.md, 避免把说明文件当作语料.
- 每个文件会被映射为一个或多个 Document, source 路径写入 metadata.
"""

from __future__ import annotations

import pathlib
import re
from typing import Any

from langchain_core.documents import Document


def _load_pdf(path: pathlib.Path) -> list[Document]:
    """加载PDF文件。"""
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except Exception as e:
        raise RuntimeError("PDF parsing requires pypdf installed") from e

    reader = PdfReader(str(path))
    docs: list[Document] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "file_type": "pdf",
                    "page": int(i + 1),
                },
            )
        )
    return docs


def _load_docx(path: pathlib.Path) -> list[Document]:
    """加载Word文档(.docx)。

    提取段落文本，并保留文档结构信息。
    """
    try:
        from docx import Document as DocxDocument  # type: ignore[import-not-found]
    except Exception as e:
        raise RuntimeError("DOCX parsing requires python-docx installed: pip install python-docx") from e

    doc = DocxDocument(str(path))

    # Extract paragraphs with heading information
    paragraphs: list[str] = []
    current_section = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if it's a heading
        if para.style.name.startswith("Heading"):
            current_section = text
            paragraphs.append(f"\n## {text}\n")
        else:
            paragraphs.append(text)

    # Also extract tables
    table_texts: list[str] = []
    for table in doc.tables:
        table_rows: list[list[str]] = []
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            table_rows.append(row_texts)

        if table_rows:
            # Convert table to markdown format
            table_str = ""
            # Header
            table_str += "| " + " | ".join(table_rows[0]) + " |\n"
            # Separator
            table_str += "|" + "|".join([" --- " for _ in table_rows[0]]) + "|\n"
            # Data rows
            for row in table_rows[1:]:
                table_str += "| " + " | ".join(row) + " |\n"
            table_texts.append(table_str)

    # Combine all content
    content = "\n".join(paragraphs)
    if table_texts:
        content += "\n\n## Tables\n\n" + "\n".join(table_texts)

    return [
        Document(
            page_content=content,
            metadata={
                "source": str(path),
                "file_type": "docx",
                "page": 0,
                "_source_text": content,
            },
        )
    ]


def _load_excel(path: pathlib.Path) -> list[Document]:
    """加载Excel文件(.xlsx, .xls)。

    将每个sheet转换为文本格式，保留表格结构。
    """
    try:
        import pandas as pd  # type: ignore[import-not-found]
    except Exception as e:
        raise RuntimeError("Excel parsing requires pandas installed: pip install pandas openpyxl xlrd") from e

    # Determine engine based on file extension
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        engine = "openpyxl"
    elif suffix == ".xls":
        engine = "xlrd"
    else:
        engine = None  # Let pandas auto-detect

    try:
        # Read all sheets
        if engine:
            sheets = pd.read_excel(str(path), sheet_name=None, engine=engine)
        else:
            sheets = pd.read_excel(str(path), sheet_name=None)
    except Exception as e:
        raise RuntimeError(f"Failed to read Excel file {path}: {e}") from e

    docs: list[Document] = []

    for sheet_name, df in sheets.items():
        # Convert DataFrame to string representation
        # Handle NaN values
        df = df.fillna("")

        # Build content with sheet name as header
        content_parts: list[str] = [f"## Sheet: {sheet_name}\n"]

        # Add column headers
        content_parts.append("Columns: " + ", ".join(str(c) for c in df.columns))
        content_parts.append("")

        # Convert rows to text
        for idx, row in df.iterrows():
            row_texts: list[str] = []
            for col, val in row.items():
                if val:
                    row_texts.append(f"{col}: {val}")
            if row_texts:
                content_parts.append(" | ".join(row_texts))

        # Also add a markdown table representation
        content_parts.append("\n### Table View\n")
        try:
            markdown_table = df.to_markdown(index=False)
            content_parts.append(str(markdown_table))
        except Exception:
            # Fallback to simple text if markdown conversion fails
            pass

        content = "\n".join(content_parts)

        docs.append(
            Document(
                page_content=content,
                metadata={
                    "source": str(path),
                    "file_type": "excel",
                    "sheet_name": str(sheet_name),
                    "page": 0,
                    "_source_text": content,
                },
            )
        )

    return docs


def _load_html(path: pathlib.Path) -> list[Document]:
    """加载HTML文件(.html, .htm)。

    提取正文内容，去除脚本和样式。
    """
    try:
        from bs4 import BeautifulSoup  # type: ignore[import-not-found]
    except Exception as e:
        raise RuntimeError("HTML parsing requires beautifulsoup4 installed: pip install beautifulsoup4 lxml") from e

    html_content = path.read_text(encoding="utf-8", errors="ignore")

    soup = BeautifulSoup(html_content, "lxml")

    # Remove script and style elements
    for script in soup(["script", "style", "nav", "footer", "header"]):
        script.decompose()

    # Extract title
    title = ""
    title_tag = soup.find("title")
    if title_tag:
        title = title_tag.get_text().strip()

    # Extract headings for structure
    headings: list[tuple[str, str]] = []
    for level in range(1, 7):
        for h in soup.find_all(f"h{level}"):
            text = h.get_text().strip()
            if text:
                headings.append((f"h{level}", text))

    # Extract main content
    # Try to find main content area
    main_content = ""

    # Look for common content containers
    for selector in ["main", "article", "[role='main']", "#content", ".content", "#main", ".main"]:
        elem = soup.select_one(selector)
        if elem:
            main_content = elem.get_text(separator="\n", strip=True)
            break

    if not main_content:
        # Fallback to body content
        body = soup.find("body")
        if body:
            main_content = body.get_text(separator="\n", strip=True)
        else:
            main_content = soup.get_text(separator="\n", strip=True)

    # Clean up excessive whitespace
    main_content = re.sub(r"\n\s*\n", "\n\n", main_content)
    main_content = re.sub(r"[ \t]+", " ", main_content)

    # Build structured content
    content_parts: list[str] = []
    if title:
        content_parts.append(f"# {title}\n")

    if headings:
        content_parts.append("## Headings\n")
        for h_level, h_text in headings[:20]:  # Limit to first 20 headings
            indent = int(h_level[1]) - 1
            content_parts.append("  " * indent + f"- {h_text}")
        content_parts.append("")

    content_parts.append("## Content\n")
    content_parts.append(main_content)

    content = "\n".join(content_parts)

    return [
        Document(
            page_content=content,
            metadata={
                "source": str(path),
                "file_type": "html",
                "page": 0,
                "title": title,
                "_source_text": content,
            },
        )
    ]


def load_sources(sources_dir: pathlib.Path) -> list[Document]:
    """从指定目录加载所有支持的文档。

    Args:
        sources_dir: 文档目录路径

    Returns:
        Document列表
    """
    # sources_dir 不存在时直接返回空列表, 方便 UI 场景下提示用户补充语料.
    if not sources_dir.exists():
        return []

    docs: list[Document] = []

    # Supported file extensions
    supported_extensions = {
        ".md": lambda p: _load_single_document(p, "md"),
        ".pdf": _load_pdf,
        ".docx": _load_docx,
        ".xlsx": _load_excel,
        ".xls": _load_excel,
        ".html": _load_html,
        ".htm": _load_html,
    }

    for path in sorted(sources_dir.rglob("*")):
        if not path.is_file():
            continue
        if path.name.lower() == "readme.md":
            continue

        suffix = path.suffix.lower().strip()

        if suffix in supported_extensions:
            try:
                loaded_docs = supported_extensions[suffix](path)
                docs.extend(loaded_docs)
            except Exception as e:
                # Log error but continue processing other files
                print(f"Warning: Failed to load {path}: {e}")
                continue

    return docs


def _load_single_document(path: pathlib.Path, file_type: str) -> list[Document]:
    """加载单个文本文件。"""
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [
        Document(
            page_content=text,
            metadata={
                "source": str(path),
                "file_type": file_type,
                "_source_text": text,
            },
        )
    ]


def get_supported_formats() -> dict[str, str]:
    """返回支持的文件格式及其描述。

    Returns:
        格式后缀到描述的映射
    """
    return {
        ".md": "Markdown文档",
        ".pdf": "PDF文档",
        ".docx": "Microsoft Word文档",
        ".xlsx": "Microsoft Excel工作表 (新格式)",
        ".xls": "Microsoft Excel工作表 (旧格式)",
        ".html": "HTML网页",
        ".htm": "HTML网页",
    }
